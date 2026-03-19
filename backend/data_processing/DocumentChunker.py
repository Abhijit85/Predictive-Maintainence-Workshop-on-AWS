from pypdf import PdfReader
import re
from typing import List, Optional, Dict, Any
import os
from pathlib import Path
from docx import Document
from bs4 import BeautifulSoup

class DocumentChunker:
    """
    Universal document chunker that supports PDF, DOCX, TXT, and HTML files.
    """
    
    def __init__(self, chunk_size: int = 1000, overlap_size: int = 200):
        """
        Initialize the document chunker with configurable parameters.
        
        Parameter:
            chunk_size: Target size for each chunk in characters
            overlap_size: Number of characters to overlap between chunks
        """
        self.chunk_size = chunk_size
        self.overlap_size = overlap_size
        self.supported_formats = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.txt': self._extract_txt,
            '.html': self._extract_html,
            '.htm': self._extract_html
        }
    
    def extract_text_from_file(self, file_path: str) -> Dict[str, Any]:
        """
        Extract text from various file formats.
        
        Parameters:
            file_path: Path to the file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        if not os.path.exists(file_path):
            return {"text": "", "error": f"File not found: {file_path}", "metadata": {}}
        
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension not in self.supported_formats:
            return {
                "text": "", 
                "error": f"Unsupported file format: {file_extension}. Supported: {list(self.supported_formats.keys())}", 
                "metadata": {}
            }
        
        try:
            extraction_func = self.supported_formats[file_extension]
            result = extraction_func(file_path)
            result["file_type"] = file_extension
            result["file_path"] = file_path
            return result
        except Exception as e:
            return {
                "text": "", 
                "error": f"Error extracting from {file_extension}: {str(e)}", 
                "metadata": {}
            }
    
    def _extract_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF file."""
        try:
            reader = PdfReader(file_path)
            text = "\n".join([page.extract_text() for page in reader.pages])
            
            metadata = {
                "pages": len(reader.pages),
                "title": reader.metadata.title if reader.metadata and reader.metadata.title else None,
                "author": reader.metadata.author if reader.metadata and reader.metadata.author else None,
                "creator": reader.metadata.creator if reader.metadata and reader.metadata.creator else None
            }
            
            return {"text": text, "error": None, "metadata": metadata}
        except Exception as e:
            return {"text": "", "error": f"PDF extraction error: {str(e)}", "metadata": {}}
    
    def _extract_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))
            
            # Combine all text
            all_text = paragraphs + table_text
            text = "\n\n".join(all_text)
            
            # Extract metadata
            metadata = {
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables),
                "title": doc.core_properties.title,
                "author": doc.core_properties.author,
                "created": doc.core_properties.created,
                "modified": doc.core_properties.modified
            }
            
            return {"text": text, "error": None, "metadata": metadata}
        except Exception as e:
            return {"text": "", "error": f"DOCX extraction error: {str(e)}", "metadata": {}}
    
    def _extract_txt(self, file_path: str) -> Dict[str, Any]:
        """Extract text from TXT file."""
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    text = file.read()
                
                # Get file stats
                file_stats = os.stat(file_path)
                metadata = {
                    "encoding": encoding,
                    "size_bytes": file_stats.st_size,
                    "lines": len(text.split('\n')),
                    "created": file_stats.st_birthtime,
                    "modified": file_stats.st_mtime
                }
                
                return {"text": text, "error": None, "metadata": metadata}
                
            except UnicodeDecodeError:
                continue
        
        return {
            "text": "", 
            "error": f"Could not decode file with encodings: {encodings}", 
            "metadata": {}
        }
    
    def _extract_html(self, file_path: str) -> Dict[str, Any]:
        """Extract text from HTML file."""
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    html_content = file.read()
                
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.extract()
                
                # Extract text
                text = soup.get_text()
                
                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = ' '.join(chunk for chunk in chunks if chunk)
                
                # Extract metadata
                title = soup.find('title')
                meta_description = soup.find('meta', attrs={'name': 'description'})
                meta_keywords = soup.find('meta', attrs={'name': 'keywords'})
                
                metadata = {
                    "encoding": encoding,
                    "title": title.string if title else None,
                    "description": meta_description.get('content') if meta_description else None,
                    "keywords": meta_keywords.get('content') if meta_keywords else None,
                    "links": len(soup.find_all('a')),
                    "images": len(soup.find_all('img')),
                    "headings": {
                        f"h{i}": len(soup.find_all(f'h{i}')) for i in range(1, 7)
                    }
                }
                
                return {"text": text, "error": None, "metadata": metadata}
                
            except UnicodeDecodeError:
                continue
        
        return {
            "text": "", 
            "error": f"Could not decode HTML file with encodings: {encodings}", 
            "metadata": {}
        }
    
    def clean_text(self, text: str) -> str:
        """
        Clean the extracted text by normalizing whitespace and removing artifacts.
        
        Parameters:
            text: Raw text to clean
            
        Returns:
            Cleaned text
        """
        # Remove extra whitespace and normalize line breaks
        text = re.sub(r'\s+', ' ', text)
        # Remove common PDF artifacts
        text = re.sub(r'[^\w\s\.\,\!\?\;\:\-\(\)\[\]\{\}\"\'\/\@\#\$\%\&\*\+\=\<\>\~\`\|\\]', '', text)
        return text.strip()
    
    def chunk_by_words(self, text: str) -> List[str]:
        """
        Chunk text using word-based overlapping strategy.
        
        Parameters:
            text: Text to chunk
            
        Returns:
            List of text chunks with word overlap
        """
        words = text.split()
        if not words:
            return []
        
        chunks = []
        start_idx = 0
        
        while start_idx < len(words):
            # Calculate end index for current chunk
            end_idx = start_idx + self._words_in_chunk(words[start_idx:])
            end_idx = min(end_idx, len(words))
            
            # Create chunk from words
            chunk = ' '.join(words[start_idx:end_idx])
            chunks.append(chunk)
            
            # If we've reached the end, break
            if end_idx >= len(words):
                break
            
            # Calculate overlap for next chunk
            overlap_words = self._words_in_overlap(words[start_idx:end_idx])
            start_idx = max(start_idx + 1, end_idx - overlap_words)
        
        return chunks
    
    def _words_in_chunk(self, words: List[str]) -> int:
        """
        Calculate how many words fit in a chunk of target character size.
        """
        char_count = 0
        word_count = 0
        
        for word in words:
            # Add word length plus space
            word_len = len(word) + 1
            if char_count + word_len > self.chunk_size and word_count > 0:
                break
            char_count += word_len
            word_count += 1
        
        return max(1, word_count)  # Ensure at least one word per chunk
    
    def _words_in_overlap(self, words: List[str]) -> int:
        """
        Calculate how many words to include in overlap.
        """
        char_count = 0
        word_count = 0
        
        # Count words from the end backwards
        for word in reversed(words):
            word_len = len(word) + 1
            if char_count + word_len > self.overlap_size:
                break
            char_count += word_len
            word_count += 1
        
        return word_count
    
    def chunk_by_sentences(self, text: str) -> List[str]:
        """
        Alternative chunking strategy using sentence boundaries with overlap.
        
        Parameters:
            text: Text to chunk
            
        Returns:
            List of text chunks with sentence-based overlap
        """
        # Split into sentences using regex
        sentence_pattern = r'(?<=[.!?])\s+'
        sentences = re.split(sentence_pattern, text)
        
        if not sentences:
            return []
        
        chunks = []
        current_chunk = ""
        overlap_sentences = []
        
        for sentence in sentences:
            # Check if adding this sentence would exceed chunk size
            potential_chunk = current_chunk + " " + sentence if current_chunk else sentence
            
            if len(potential_chunk) > self.chunk_size and current_chunk:
                # Finalize current chunk
                chunks.append(current_chunk.strip())
                
                # Start new chunk with overlap sentences
                current_chunk = " ".join(overlap_sentences) + " " + sentence
                overlap_sentences = []
            else:
                current_chunk = potential_chunk
            
            # Maintain overlap sentences (last few sentences that fit in overlap_size)
            temp_overlap = overlap_sentences + [sentence]
            overlap_text = " ".join(temp_overlap)
            
            while len(overlap_text) > self.overlap_size and len(temp_overlap) > 1:
                temp_overlap.pop(0)
                overlap_text = " ".join(temp_overlap)
            
            overlap_sentences = temp_overlap
        
        # Add the last chunk if it exists
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def chunk_by_paragraphs(self, text: str) -> List[str]:
        """
        Chunking strategy using paragraph boundaries with overlap.
        This method preserves paragraph structure while maintaining target chunk sizes.
        
        Parameters:
            text: Text to chunk
            
        Returns:
            List of text chunks with paragraph-based overlap
        """
        # Split into paragraphs - handle different paragraph separators
        # Common patterns: double newlines, single newlines with indentation, etc.
        paragraphs = self._split_into_paragraphs(text)
        
        if not paragraphs:
            return []
        
        chunks = []
        current_chunk = ""
        overlap_paragraphs = []
        
        for paragraph in paragraphs:
            # Skip empty paragraphs
            if not paragraph.strip():
                continue
            
            # Check if adding this paragraph would exceed chunk size
            potential_chunk = self._join_paragraphs([current_chunk, paragraph]) if current_chunk else paragraph
            
            if len(potential_chunk) > self.chunk_size and current_chunk:
                # Finalize current chunk
                chunks.append(current_chunk.strip())
                
                # Start new chunk with overlap paragraphs plus current paragraph
                if overlap_paragraphs:
                    current_chunk = self._join_paragraphs(overlap_paragraphs + [paragraph])
                else:
                    current_chunk = paragraph
                overlap_paragraphs = []
            else:
                current_chunk = potential_chunk
            
            # Maintain overlap paragraphs (last few paragraphs that fit in overlap_size)
            temp_overlap = overlap_paragraphs + [paragraph]
            overlap_text = self._join_paragraphs(temp_overlap)
            
            # Remove paragraphs from beginning of overlap if it exceeds overlap_size
            while len(overlap_text) > self.overlap_size and len(temp_overlap) > 1:
                temp_overlap.pop(0)
                overlap_text = self._join_paragraphs(temp_overlap)
            
            overlap_paragraphs = temp_overlap
        
        # Add the last chunk if it exists
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks
    
    def _split_into_paragraphs(self, text: str) -> List[str]:
        """
        Split text into paragraphs using various heuristics.
        
        Parameters:
            text: Text to split
            
        Returns:
            List of paragraphs
        """
        # Method 1: Split on double newlines (most common)
        paragraphs = re.split(r'\n\s*\n', text)
        
        # If we don't get good paragraph separation, try other methods
        if len(paragraphs) < 3:
            # Method 2: Split on single newlines followed by capitalized words (common in PDFs)
            paragraphs = re.split(r'\n(?=[A-Z][a-z])', text)
        
        # If still not good separation, try detecting indented paragraphs
        if len(paragraphs) < 3:
            # Method 3: Split on newlines followed by spaces (indented paragraphs)
            paragraphs = re.split(r'\n(?=\s{2,})', text)
        
        # Clean up paragraphs
        cleaned_paragraphs = []
        for para in paragraphs:
            # Remove extra whitespace and normalize
            para = re.sub(r'\s+', ' ', para.strip())
            if para and len(para) > 10:  # Filter out very short paragraphs
                cleaned_paragraphs.append(para)
        
        # If we still don't have good paragraphs, fall back to sentence-based splitting
        if len(cleaned_paragraphs) < 2:
            # Split into pseudo-paragraphs based on sentence groups
            sentences = re.split(r'(?<=[.!?])\s+', text)
            paragraphs = []
            current_para = ""
            
            for sentence in sentences:
                if len(current_para + " " + sentence) > 300:  # Arbitrary paragraph size
                    if current_para:
                        paragraphs.append(current_para.strip())
                    current_para = sentence
                else:
                    current_para = current_para + " " + sentence if current_para else sentence
            
            if current_para:
                paragraphs.append(current_para.strip())
            
            return paragraphs
        
        return cleaned_paragraphs
    
    def _join_paragraphs(self, paragraphs: List[str]) -> str:
        """
        Join paragraphs with appropriate separators.
        
        Parameters:
            paragraphs: List of paragraph strings
            
        Returns:
            Joined text
        """
        # Filter out empty paragraphs
        valid_paragraphs = [p for p in paragraphs if p and p.strip()]
        
        # Join with double newlines to preserve paragraph structure
        return '\n\n'.join(valid_paragraphs)
    
    def process_document(self, file_path: str, method: str = "words") -> Dict[str, Any]:
        """
        Complete pipeline to extract and chunk document content from various formats.
        
        Parameters:
            file_path: Path to the document file
            method: Chunking method ("words", "sentences", or "paragraphs")
            
        Returns:
            Dictionary containing chunks, metadata, and processing info
        """
        # Extract text and metadata
        extraction_result = self.extract_text_from_file(file_path)
        
        if extraction_result["error"]:
            return {
                "chunks": [],
                "error": extraction_result["error"],
                "metadata": extraction_result["metadata"],
                "file_info": {"path": file_path, "type": Path(file_path).suffix.lower()}
            }
        
        raw_text = extraction_result["text"]
        if not raw_text:
            return {
                "chunks": [],
                "error": "No text content found in file",
                "metadata": extraction_result["metadata"],
                "file_info": {"path": file_path, "type": Path(file_path).suffix.lower()}
            }
        
        # Clean text
        cleaned_text = self.clean_text(raw_text)
        
        # Chunk text based on method
        if method == "sentences":
            chunks = self.chunk_by_sentences(cleaned_text)
        elif method == "paragraphs":
            chunks = self.chunk_by_paragraphs(cleaned_text)
        else:
            chunks = self.chunk_by_words(cleaned_text)
        
        return {
            "chunks": chunks,
            "error": None,
            "metadata": extraction_result["metadata"],
            "file_info": {
                "path": file_path,
                "type": Path(file_path).suffix.lower(),
                "original_length": len(raw_text),
                "cleaned_length": len(cleaned_text),
                "total_chunks": len(chunks)
            }
        }

    def process_multiple_files(self, file_paths: List[str], method: str = "words") -> Dict[str, Dict[str, Any]]:
        """
        Process multiple files and return results for each.
        
        Parameters:
            file_paths: List of file paths to process
            method: Chunking method to use
            
        Returns:
            Dictionary with file paths as keys and processing results as values
        """
        results = {}
        
        for file_path in file_paths:
            print(f"Processing: {file_path}")
            result = self.process_document(file_path, method)
            results[file_path] = result
            
            if result["error"]:
                print(f"  Error: {result['error']}")
            else:
                print(f"  Success: {len(result['chunks'])} chunks created")
        
        return results

    def print_chunk_stats(self, chunks: List[str], metadata: Dict[str, Any] = None) -> None:
        """
        Print statistics about the chunks and document metadata.
        """
        if not chunks:
            print("No chunks generated.")
            return
        
        chunk_lengths = [len(chunk) for chunk in chunks]
        
        print(f"Total chunks: {len(chunks)}")
        print(f"Average chunk length: {sum(chunk_lengths) / len(chunk_lengths):.0f} characters")
        print(f"Min chunk length: {min(chunk_lengths)} characters")
        print(f"Max chunk length: {max(chunk_lengths)} characters")
        
        # Print metadata if available
        if metadata:
            print(f"\nDocument Metadata:")
            for key, value in metadata.items():
                if value is not None and value != "":
                    print(f"  {key}: {value}")
